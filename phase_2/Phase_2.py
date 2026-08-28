import pymupdf
import docx

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pymupdf.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    doc = docx.Document(file_path)
    
    # 1. Đọc text từ các đoạn văn bình thường
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"
            
    # 2. Đọc text từ các Bảng biểu (Tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text.strip() + "\n"
                    
    return text

def extract_text(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")
    
print(extract_text("D:\\G.VAN\\ai_foundation\\phase_2\\doc.docx"))