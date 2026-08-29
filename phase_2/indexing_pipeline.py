import pymupdf
import docx
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

# ─── Config ────────────────────────────────────────────────────────────
COLLECTION_NAME = "my_docs"
MODEL_NAME      = "all-MiniLM-L6-v2"
VECTOR_SIZE     = 384
CHUNK_SIZE      = 500
CHUNK_OVERLAP   = 50

# ─── Core Functions ────────────────────────────────────────────────────
def extract_text(file_path: str) -> str:
    """Đọc text từ file PDF hoặc DOCX."""
    text = ""
    if file_path.endswith(".pdf"):
        with pymupdf.open(file_path) as doc:
            text = "".join(page.get_text() for page in doc)
            
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        # Gom text từ paragraphs
        text += "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip()) + "\n"
        # Gom text từ bảng biểu
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text.strip() + "\n"
    else:
        raise ValueError(f"Không hỗ trợ định dạng file: {file_path}")
    return text
 
def split_text(text: str) -> list[str]:
    """Tách text thành các chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    return splitter.split_text(text)

# ─── Pipeline ──────────────────────────────────────────────────────────
def index_document(file_path: str, model: SentenceTransformer, client: QdrantClient):
    """Toàn bộ quy trình: Extract -> Chunk -> Embed -> Upsert."""
    print(f"\n[1] Đang xử lý file: {file_path}")
    
    # 1. Trích xuất & Chunking
    text = extract_text(file_path)
    chunks = split_text(text)
    print(f"  → Đã tạo {len(chunks)} chunks.")

    # 2. Embedding
    print("[2] Đang tính toán embedding...")
    vectors = model.encode(chunks, batch_size=32, show_progress_bar=True)

    # 3. Tạo/Reset Collection trong Qdrant
    client.recreate_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=VECTOR_SIZE, distance=Distance.COSINE),
    )

    # 4. Upsert (Lưu vào Vector DB)
    print("[3] Đang lưu vào Qdrant...")
    points = [
        PointStruct(
            id=i,
            vector=vec.tolist(),
            payload={"text": chunk, "source": file_path, "chunk_id": i},
        )
        for i, (chunk, vec) in enumerate(zip(chunks, vectors))
    ]
    client.upsert(collection_name=COLLECTION_NAME, points=points)
    print(f"  → Đã index thành công {len(points)} chunks vào '{COLLECTION_NAME}'.")

def search(query: str, model: SentenceTransformer, client: QdrantClient, top_k: int = 3):
    """Tìm kiếm semantic search."""
    q_vec = model.encode(query).tolist()
    # Ở phiên bản qdrant-client mới, hàm search được thay bằng query_points
    response = client.query_points(collection_name=COLLECTION_NAME, query=q_vec, limit=top_k)
    
    print(f"\n🔍 Query: '{query}'")
    for i, r in enumerate(response.points, 1):
        print(f"  {i}. [score={r.score:.4f}] {r.payload['text'][:150].replace('\n', ' ')}...")


# ─── Main Execution ────────────────────────────────────────────────────
if __name__ == "__main__":
    # Khởi tạo Heavy Objects 1 lần duy nhất
    print("Đang tải model và kết nối Qdrant...")
    model = SentenceTransformer(MODEL_NAME)
    client = QdrantClient(host="localhost", port=6333)

    # 1. Chạy index file (thay bằng đường dẫn file thật)
    file_to_index = "resources/doc.docx"  # Dùng dấu / để tránh lỗi escape character
    if Path(file_to_index).exists():
        index_document(file_to_index, model, client)
        # 2. Test search (Chỉ search nếu đã index thành công hoặc collection đã có)
        search("When will we introduce Data Analytics? Python Review? D3M Process?", model, client)
    else:
        print(f"⚠️ File không tồn tại: {file_to_index}. Đảm bảo bạn đang chạy script trong thư mục phase_2.")
